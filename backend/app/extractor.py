"""
Enhanced Document Extraction Module
Phase 5: Multi-format document extraction using Docling, PyPDF2, pdfplumber, and python-docx
"""

from typing import Dict, Any, Optional, List
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentExtractor:
    """
    Unified document extraction interface supporting multiple formats
    Primary: Docling for advanced extraction
    Fallback: Format-specific libraries
    """

    def __init__(self):
        self.docling_available = False
        self.pypdf2_available = False
        self.pdfplumber_available = False
        self.docx_available = False

        # Try to import Docling
        try:
            from docling.document_converter import DocumentConverter
            self.DocumentConverter = DocumentConverter
            self.docling_available = True
            logger.info("Docling loaded successfully")
        except ImportError as e:
            logger.warning(f"Docling not available: {e}")

        # Try to import PyPDF2
        try:
            import PyPDF2
            self.PyPDF2 = PyPDF2
            self.pypdf2_available = True
            logger.info("PyPDF2 loaded successfully")
        except ImportError as e:
            logger.warning(f"PyPDF2 not available: {e}")

        # Try to import pdfplumber
        try:
            import pdfplumber
            self.pdfplumber = pdfplumber
            self.pdfplumber_available = True
            logger.info("pdfplumber loaded successfully")
        except ImportError as e:
            logger.warning(f"pdfplumber not available: {e}")

        # Try to import python-docx
        try:
            from docx import Document
            self.DocxDocument = Document
            self.docx_available = True
            logger.info("python-docx loaded successfully")
        except ImportError as e:
            logger.warning(f"python-docx not available: {e}")

    def extract_text(
        self,
        file_path: str,
        file_type: str,
        use_docling: bool = True
    ) -> Dict[str, Any]:
        """
        Extract text from document with automatic format detection

        Args:
            file_path: Path to document file
            file_type: File type (txt, pdf, docx, md)
            use_docling: Try Docling first if available (recommended)

        Returns:
            Dict with extracted text, metadata, and extraction method
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # TXT and MD files - simple text extraction
        if file_type in ["txt", "md"]:
            return self._extract_text_simple(file_path)

        # PDF files
        elif file_type == "pdf":
            return self._extract_pdf(file_path, use_docling)

        # DOCX files
        elif file_type == "docx":
            return self._extract_docx(file_path, use_docling)

        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_text_simple(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT/MD files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()

        return {
            "text": text,
            "method": "simple_text",
            "pages": 1,
            "has_tables": False,
            "has_images": False,
            "metadata": {}
        }

    def _extract_pdf(self, file_path: Path, use_docling: bool) -> Dict[str, Any]:
        """
        Extract text from PDF with multiple fallback strategies
        1. Try Docling (best quality, handles layout, tables, images)
        2. Try pdfplumber (good for tables)
        3. Fallback to PyPDF2 (basic text extraction)
        """
        # Strategy 1: Docling (preferred)
        if use_docling and self.docling_available:
            try:
                return self._extract_pdf_docling(file_path)
            except Exception as e:
                logger.warning(f"Docling extraction failed: {e}, trying fallback methods")

        # Strategy 2: pdfplumber (good for tables)
        if self.pdfplumber_available:
            try:
                return self._extract_pdf_pdfplumber(file_path)
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}, trying PyPDF2")

        # Strategy 3: PyPDF2 (basic fallback)
        if self.pypdf2_available:
            try:
                return self._extract_pdf_pypdf2(file_path)
            except Exception as e:
                logger.error(f"PyPDF2 extraction failed: {e}")
                raise

        raise RuntimeError("No PDF extraction library available")

    def _extract_pdf_docling(self, file_path: Path) -> Dict[str, Any]:
        """Extract PDF using Docling (advanced)"""
        converter = self.DocumentConverter()
        result = converter.convert(str(file_path))

        # Export to markdown format
        markdown_text = result.document.export_to_markdown()

        # Get metadata
        metadata = {
            "title": getattr(result.document, "title", ""),
            "num_pages": len(getattr(result.document, "pages", [])),
        }

        return {
            "text": markdown_text,
            "method": "docling",
            "pages": metadata.get("num_pages", 1),
            "has_tables": True,  # Docling handles tables
            "has_images": True,  # Docling can extract images
            "metadata": metadata
        }

    def _extract_pdf_pdfplumber(self, file_path: Path) -> Dict[str, Any]:
        """Extract PDF using pdfplumber (good for tables)"""
        text_parts = []
        tables = []

        with self.pdfplumber.open(file_path) as pdf:
            num_pages = len(pdf.pages)

            for page in pdf.pages:
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

                # Extract tables
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)

        # Combine text
        text = "\n\n".join(text_parts)

        # Add tables as formatted text
        if tables:
            text += "\n\n--- Tables ---\n\n"
            for i, table in enumerate(tables):
                text += f"\nTable {i+1}:\n"
                for row in table:
                    text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"

        return {
            "text": text,
            "method": "pdfplumber",
            "pages": num_pages,
            "has_tables": len(tables) > 0,
            "has_images": False,
            "metadata": {
                "num_tables": len(tables)
            }
        }

    def _extract_pdf_pypdf2(self, file_path: Path) -> Dict[str, Any]:
        """Extract PDF using PyPDF2 (basic fallback)"""
        text_parts = []

        with open(file_path, 'rb') as f:
            pdf_reader = self.PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        text = "\n\n".join(text_parts)

        return {
            "text": text,
            "method": "pypdf2",
            "pages": num_pages,
            "has_tables": False,
            "has_images": False,
            "metadata": {}
        }

    def _extract_docx(self, file_path: Path, use_docling: bool) -> Dict[str, Any]:
        """
        Extract text from DOCX
        1. Try Docling (preserves formatting)
        2. Fallback to python-docx
        """
        # Strategy 1: Docling (preferred)
        if use_docling and self.docling_available:
            try:
                return self._extract_docx_docling(file_path)
            except Exception as e:
                logger.warning(f"Docling DOCX extraction failed: {e}, trying python-docx")

        # Strategy 2: python-docx (fallback)
        if self.docx_available:
            try:
                return self._extract_docx_python_docx(file_path)
            except Exception as e:
                logger.error(f"python-docx extraction failed: {e}")
                raise

        raise RuntimeError("No DOCX extraction library available")

    def _extract_docx_docling(self, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX using Docling"""
        converter = self.DocumentConverter()
        result = converter.convert(str(file_path))

        # Export to markdown format
        markdown_text = result.document.export_to_markdown()

        return {
            "text": markdown_text,
            "method": "docling",
            "pages": 1,
            "has_tables": True,
            "has_images": True,
            "metadata": {}
        }

    def _extract_docx_python_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX using python-docx"""
        doc = self.DocxDocument(str(file_path))

        # Extract paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract tables
        tables = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                table_text.append(row_text)
            tables.append("\n".join(table_text))

        # Combine text and tables
        text = "\n\n".join(text_parts)

        if tables:
            text += "\n\n--- Tables ---\n\n"
            text += "\n\n".join(tables)

        return {
            "text": text,
            "method": "python-docx",
            "pages": 1,
            "has_tables": len(tables) > 0,
            "has_images": False,
            "metadata": {
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables)
            }
        }

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats based on available libraries"""
        formats = ["txt", "md"]  # Always supported

        if self.docling_available or self.pypdf2_available or self.pdfplumber_available:
            formats.append("pdf")

        if self.docling_available or self.docx_available:
            formats.append("docx")

        return formats

    def get_extraction_capabilities(self) -> Dict[str, Any]:
        """Get information about available extraction methods"""
        return {
            "docling": self.docling_available,
            "pypdf2": self.pypdf2_available,
            "pdfplumber": self.pdfplumber_available,
            "python_docx": self.docx_available,
            "supported_formats": self.get_supported_formats()
        }


# Global extractor instance
extractor = DocumentExtractor()
